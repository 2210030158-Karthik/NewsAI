from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading_with_body(doc: Document, heading: str, body: str) -> None:
    doc.add_heading(heading, level=1)
    for paragraph in body.strip().split("\n\n"):
        doc.add_paragraph(paragraph.strip())


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_v1_v2_table(doc: Document) -> None:
    doc.add_heading("V1 to V2 Comparison Summary", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "Version 1"
    hdr[2].text = "Version 2"

    rows = [
        (
            "Model Execution",
            "Primary NLP and ranking behavior depended on external API calls.",
            "Core topic classification and ranking now run locally with fallback logic, reducing dependency on external model APIs.",
        ),
        (
            "Ingestion Mode",
            "Request-driven and partially synchronous fetch flow.",
            "Asynchronous queue-based ingestion using Celery + Redis with scheduled 5-hour sweep support.",
        ),
        (
            "Scalability",
            "Not designed for large topic batches in production-like cadence.",
            "Supports high-volume active-topic sweeps (up to 250-topic batches) and progress persistence checkpoints during runs.",
        ),
        (
            "Article Experience",
            "Feed focused mainly on headline-level metadata and source redirect.",
            "In-app full-content reader with extraction refresh endpoint and fallback handling for blocked publishers.",
        ),
        (
            "Ranking",
            "Basic recency/match ordering.",
            "Freshness-first plus relevance ranking with feedback-aware topic and source weighting.",
        ),
        (
            "Feedback Learning",
            "Limited closed-loop personalization.",
            "Like/dislike feedback updates user preference profiles and influences subsequent feed scoring.",
        ),
        (
            "Reliability",
            "Frequent runtime confusion from stale listeners and mixed service state.",
            "One-command orchestration, improved process cleanup, port strategy updates, and clearer run observability.",
        ),
        (
            "Observability",
            "Hard to distinguish true errors from noise.",
            "Run metrics, ingestion tracking, and cleaner logging with explicit run/task summaries.",
        ),
    ]

    for area, v1, v2 in rows:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = v1
        cells[2].text = v2


def main() -> None:
    doc = Document()

    title = doc.add_paragraph()
    title_run = title.add_run("Personalized News Platform\nFinal Year Project Report (Version 2)")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        "Prepared for: Final Year Project Evaluation\n"
        "Project: AI-Powered Personalized News Intelligence Platform\n"
        f"Report Date: {datetime.now().strftime('%d %B %Y')}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph(
        "Note: This report documents the upgraded V2 system, highlights the engineering evolution from V1, "
        "and provides implementation-level explanations suitable for academic review."
    )

    add_heading_with_body(
        doc,
        "Abstract",
        "This report presents Version 2 of the Personalized News Platform, an AI-driven system that ingests, ranks, and serves topic-relevant news for individual users. "
        "Compared to Version 1, the V2 architecture introduces asynchronous large-batch ingestion, local model inference for core classification and ranking operations, robust full-article extraction, and stronger runtime reliability.\n\n"
        "A central design improvement is reduced dependence on external APIs for AI behavior. In V1, model behavior was largely API-centered. In V2, essential model processing paths are executed locally, with resilient fallback mechanisms. "
        "This change improves cost predictability, latency stability, and operational control while preserving system flexibility.\n\n"
        "The resulting platform is better aligned with production-grade requirements: continuous ingestion, feedback-aware personalization, and a clearer, more maintainable service boundary between API layer, worker layer, and user-facing experience."
    )

    add_heading_with_body(
        doc,
        "Introduction",
        "The information ecosystem is dominated by high-velocity, multi-source content that can overwhelm users and reduce decision quality. A personalized news platform must therefore do more than aggregate headlines; it should continuously learn user interest patterns, rank fresh information effectively, and present content in a readable format.\n\n"
        "Version 1 established the baseline product concept, including user authentication, topic-based feed generation, and initial article retrieval. Version 2 advances this baseline into a more mature engineering system by introducing asynchronous ingestion, local model-first intelligence workflows, robust article extraction, and better operational observability."
    )

    add_heading_with_body(
        doc,
        "Project Objectives",
        "The primary objectives of Version 2 were to improve system capability, reliability, and academic/technical quality. The enhancement goals were:"
    )
    add_bullet_list(
        doc,
        [
            "Deliver scalable ingestion capable of processing large active-topic batches on a schedule.",
            "Shift core AI inference away from all-API dependence toward local model execution.",
            "Enable full in-app article reading instead of forcing users to open external sources for every story.",
            "Introduce feedback-aware ranking so user actions (like/dislike) influence future feed quality.",
            "Improve runtime reliability for local development on Windows through stronger process orchestration.",
            "Provide measurable ingestion progress and cleaner operational telemetry during long runs.",
        ],
    )

    add_heading_with_body(
        doc,
        "System Overview (Version 1 Baseline)",
        "Version 1 implemented a functional but limited end-to-end pipeline. The platform could authenticate users, map users to selected topics, fetch external article metadata, and show a personalized feed. However, it faced practical limitations in scale and robustness.\n\n"
        "The most important limitation was architectural coupling between external dependencies and core product behavior. Model intelligence and parts of ranking behavior were heavily API-oriented, reducing control over performance and failure modes. "
        "Ingestion execution was also less resilient under sustained batch workloads, and article readability often depended on source-site redirects."
    )

    add_heading_with_body(
        doc,
        "Enhanced System Design (Version 2)",
        "Version 2 introduces a modular architecture with clearer service boundaries. The API layer handles user interactions and orchestration endpoints. Background workers execute ingestion, extraction, and classification tasks asynchronously. Redis-backed queuing enables non-blocking execution and scheduled workloads.\n\n"
        "The ingestion pipeline follows a structured flow: discovery -> fetch -> extract -> normalize -> classify -> persist -> rank. This flow supports better fault isolation and enables checkpoint-style progress updates during larger runs.\n\n"
        "At the user level, the feed now reflects freshness and relevance together, while the in-app reader and hydration endpoint improve content accessibility without requiring source-site redirection for every interaction."
    )

    add_heading_with_body(
        doc,
        "Local Model Strategy vs API-Only Dependence",
        "A major V2 improvement is the shift from broad API dependence to local model execution for core intelligence paths.\n\n"
        "In Version 1, model-driven functions were mainly executed through external calls, which made quality and latency dependent on external service conditions. In Version 2, local classifier controls and model selection are introduced, including support for smaller local models and robust lexical fallback behavior. "
        "This ensures that the platform can continue making relevance decisions even when external model services are unavailable or expensive.\n\n"
        "External APIs are still used where they are strategically valuable (for example, live discovery data sources), but not for everything. Core personalization logic is now owned inside the application stack. "
        "This hybrid model is intentionally designed to maximize control, resilience, and maintainability."
    )

    add_heading_with_body(
        doc,
        "Implementation Highlights in Version 2",
        "Version 2 development focused on practical engineering outcomes, not just feature additions. Key implementation upgrades include the following:"
    )
    add_bullet_list(
        doc,
        [
            "Asynchronous ingestion workflow using Celery + Redis with queue-backed processing.",
            "Scheduled active-topic sweep support for periodic large-batch ingestion.",
            "Progress checkpoint persistence during runs (every 20 processed article URLs) for live DB visibility.",
            "Full-content retrieval endpoint with hydration for articles missing sufficient extracted text.",
            "Feedback learning loop that updates per-user topic and source preference profiles.",
            "Freshness-first and relevance-aware feed ordering to surface newly fetched content clearly.",
            "Local model loading controls with safe fallback behavior.",
            "Launcher/orchestration improvements for stable multi-service startup and recovery on Windows.",
        ],
    )

    add_v1_v2_table(doc)

    add_heading_with_body(
        doc,
        "Results and Observations",
        "Version 2 demonstrates stronger runtime behavior and better user-facing utility than Version 1. Ingestion runs complete with explicit run metrics, and the platform can process large topic sweeps without blocking API requests. "
        "Feed quality benefits from both freshness and personalization signals, while full-content availability has improved through extract-and-hydrate logic.\n\n"
        "Operationally, the project now has clearer diagnostic insight. Instead of relying on ambiguous console output, ingestion status and summary metrics can be inspected directly through API endpoints and database records."
    )

    add_heading_with_body(
        doc,
        "Professional Reflection and Engineering Rationale",
        "The transition from V1 to V2 reflects a maturity shift from prototype behavior to engineered platform behavior. The design intentionally avoids over-reliance on third-party APIs for critical logic. "
        "By moving key model workflows locally, the system gains deterministic control over core product quality and reduces external failure impact.\n\n"
        "From a software engineering perspective, this transition improves reproducibility, testability, and long-term maintainability. It also better aligns with final-year project evaluation criteria that emphasize system design decisions, trade-off analysis, and technical depth."
    )

    add_heading_with_body(
        doc,
        "Conclusion",
        "Version 2 successfully extends the Personalized News Platform into a scalable, resilient, and academically robust system. The architecture now supports asynchronous processing at higher topic volumes, local model-first personalization logic, improved article readability, and clearer operational insight.\n\n"
        "Most importantly, the project now demonstrates a balanced hybrid strategy: external APIs are used selectively for data discovery, while core intelligence and ranking behavior are managed locally. "
        "This is a significant technical improvement over Version 1 and represents a more production-aligned solution for personalized news delivery."
    )

    add_heading_with_body(
        doc,
        "Future Work",
        "Potential future improvements include:"
    )
    add_bullet_list(
        doc,
        [
            "Advanced source-quality scoring and adaptive crawl prioritization.",
            "Richer explainability panels showing why each article was ranked.",
            "Automated stale-run recovery and stronger operational dashboards.",
            "Model benchmarking framework for continuous local-model evaluation.",
            "User-configurable freshness windows and ranking strategy presets.",
        ],
    )

    add_heading_with_body(
        doc,
        "Screenshot Placement Guide",
        "You can insert screenshots under the following sections for maximum report clarity:\n\n"
        "1) Architecture and pipeline flow diagram under Enhanced System Design.\n"
        "2) Queue run metrics and ingestion status UI under Implementation Highlights.\n"
        "3) Feed interface and in-app reader under Results and Observations.\n"
        "4) V1 vs V2 evidence snapshots (logs, endpoints, model configuration) under Comparison Summary."
    )

    output_path = r"d:\Projects\personalized-news\Personalized_News_Platform_V2_Report.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
